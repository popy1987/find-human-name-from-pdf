import logging
import os
import subprocess
import sys
from pathlib import Path
from typing import List, Optional


class Prepare:
    """Prepare class for setup and initialization tasks.

    功能职责：
    1. 检查并安装必要的 spaCy 语言模型（优先从本地 dic 目录加载）
    2. 校验文件（路径、格式、可读性）
    3. 根据文件类型加载对应的解析器（目前仅支持 PDF）
    4. 将文件内容读取为文本（处理大文件，支持几十万字）
    """

    SUPPORTED_FORMATS = {".pdf", ".doc", ".docx", ".txt"}
    REQUIRED_SPACY_MODELS = ["zh_core_web_sm", "en_core_web_sm"]

    def __init__(self, file_path: str, logger: logging.Logger):
        """Initialize the Prepare class.

        Args:
            file_path: 输入文件的绝对路径
            logger: 日志记录器实例
        """
        self.file_path = file_path
        self.logger = logger
        self.file_ext: Optional[str] = None

        # 本地模型目录
        self.dic_dir = os.path.join(os.path.dirname(__file__), "dic")
        self._ensure_dic_dir()

        # 用户自定义人名词典路径
        self.name_dic_path = os.path.join(self.dic_dir, "name_dic_for_user.txt")

        # 加载的自定义人名列表（供 Run 使用）
        self.custom_names: List[str] = []

    def _ensure_dic_dir(self):
        """确保本地模型目录存在。"""
        os.makedirs(self.dic_dir, exist_ok=True)

    def _get_local_model_path(self, model_name: str) -> Optional[str]:
        """获取本地模型路径，兼容多种目录结构。

        支持的目录结构：
        1. dic/{model}/config.cfg  - 直接放置的模型
        2. dic/{model}/{model}-3.x.x/config.cfg  - pip 安装后的包结构（子目录）

        Returns:
            如果本地存在可加载的模型则返回完整路径，否则返回 None
        """
        base_path = os.path.join(self.dic_dir, model_name)
        if not os.path.exists(base_path) or not os.path.isdir(base_path):
            return None

        # 结构1: config.cfg 直接在模型目录下
        config_direct = os.path.join(base_path, "config.cfg")
        if os.path.exists(config_direct):
            return base_path

        # 结构2: 模型在子目录 {model_name}-版本号/ 下（pip 安装的包结构）
        try:
            for item in os.listdir(base_path):
                sub_path = os.path.join(base_path, item)
                if os.path.isdir(sub_path) and item.startswith(model_name + "-"):
                    config_in_sub = os.path.join(sub_path, "config.cfg")
                    if os.path.exists(config_in_sub):
                        return sub_path
        except OSError:
            pass

        return None

    def _check_spacy_model(self, model_name: str) -> tuple[bool, Optional[str]]:
        """检查指定的 spaCy 模型是否可用。

        优先顺序：1. 本地 dic 目录 -> 2. 系统已安装

        Returns:
            (是否可用, 路径或名称)
        """
        import spacy

        # 1. 优先检查本地目录
        local_path = self._get_local_model_path(model_name)
        if local_path:
            try:
                spacy.load(local_path)
                return True, local_path
            except Exception:
                pass

        # 2. 检查系统已安装
        try:
            spacy.load(model_name)
            return True, model_name
        except (ImportError, OSError):
            return False, None

    def _install_model_to_local(self, model_name: str) -> bool:
        """下载并安装模型到本地 dic 目录。

        Returns:
            True 如果成功安装到本地
        """
        local_path = os.path.join(self.dic_dir, model_name)

        self.logger.info(f"正在下载 spaCy 模型 {model_name} 到本地目录...")
        self.logger.info(f"目标路径: {local_path}")

        try:
            # 使用 spacy download 命令（--target 为 pip 参数，指定安装目录）
            result = subprocess.run(
                [sys.executable, "-m", "spacy", "download", model_name, "--target", self.dic_dir],
                capture_output=True,
                text=True,
                check=True,
            )

            # 验证安装
            if self._get_local_model_path(model_name):
                self.logger.info(f"✓ 模型 {model_name} 已成功安装到本地目录")
                return True
            else:
                self.logger.error(f"模型下载后未在预期位置找到")
                return False

        except subprocess.CalledProcessError as e:
            self.logger.error(f"安装模型 {model_name} 失败: {e}")
            self.logger.error(f"错误输出: {e.stderr}")
            return False

    def validate(self) -> bool:
        """第一步：校验文件。

        校验项：
        - 文件是否存在
        - 是否为文件（非目录）
        - 文件大小是否为 0
        - 文件扩展名是否支持
        - 文件是否可读

        Returns:
            True 如果校验通过，否则抛出异常
        """
        self.logger.info(f"开始校验文件: {self.file_path}")

        path = Path(self.file_path)

        # 1. 检查文件是否存在
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {self.file_path}")

        # 2. 检查是否为文件
        if not path.is_file():
            raise ValueError(f"路径不是文件: {self.file_path}")

        # 3. 检查文件大小
        file_size = path.stat().st_size
        if file_size == 0:
            raise ValueError(f"文件为空: {self.file_path}")
        self.logger.info(f"文件大小: {file_size} bytes")

        # 4. 检查文件扩展名
        self.file_ext = path.suffix.lower()
        if self.file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"不支持的文件格式: {self.file_ext}. "
                f"目前支持: {', '.join(self.SUPPORTED_FORMATS)}"
            )
        self.logger.info(f"文件格式校验通过: {self.file_ext}")

        # 5. 检查文件是否可读
        if not os.access(self.file_path, os.R_OK):
            raise PermissionError(f"文件不可读: {self.file_path}")

        self.logger.info("文件校验通过")
        return True

    def ensure_spacy_models(self) -> bool:
        """第一步：确保必需的 spaCy 模型都在 dic 目录下。

        校验规则：若 dic 下缺少 zh_core_web_sm 或 en_core_web_sm，
        则自动执行下载命令安装到 dic 目录。不依赖系统已安装的模型。

        Returns:
            True 如果所有模型都在 dic 中就绪

        Raises:
            RuntimeError: 如果模型下载失败
        """
        self.logger.info("=== 检查 spaCy 语言模型 ===")
        self.logger.info(f"本地模型目录: {self.dic_dir}")

        try:
            import spacy
            self.logger.info(f"spaCy 版本: {spacy.__version__}")
        except ImportError:
            raise ImportError(
                "未安装 spaCy，请先运行: pip install spacy"
            )

        missing_in_dic = [
            m for m in self.REQUIRED_SPACY_MODELS
            if self._get_local_model_path(m) is None
        ]

        if missing_in_dic:
            self.logger.info(f"dic 下缺少模型: {', '.join(missing_in_dic)}，开始下载到 dic...")
            for model in missing_in_dic:
                self.logger.info(f"正在下载 {model} 到 {self.dic_dir}...")
                if not self._install_model_to_local(model):
                    raise RuntimeError(
                        f"模型 {model} 下载失败。请手动运行: "
                        f"python -m spacy download {model} --target {self.dic_dir}"
                    )
                self.logger.info(f"✓ {model} 已安装到 dic")
        else:
            for model in self.REQUIRED_SPACY_MODELS:
                path = self._get_local_model_path(model)
                self.logger.info(f"✓ 模型 {model} 已在 dic 中: {path}")

        self.logger.info("所有 spaCy 模型准备就绪")
        return True

    def _load_custom_names(self) -> List[str]:
        """加载 name_dic_for_user.txt 中的自定义人名。

        若文件不存在则创建模板并返回空列表。
        每行一个名字，以 # 开头的行和空行会被忽略。

        Returns:
            自定义人名的列表
        """
        if not os.path.exists(self.name_dic_path):
            default_content = (
                "# 用户自定义人名词典\n"
                "# 每行一个名字，支持中文和英文\n"
                "# 以 # 开头的行会被忽略\n"
            )
            try:
                with open(self.name_dic_path, "w", encoding="utf-8") as f:
                    f.write(default_content)
                self.logger.info(f"已创建模板文件: {self.name_dic_path}")
            except OSError as e:
                self.logger.warning(f"无法创建 name_dic_for_user.txt: {e}")
            return []

        names = []
        try:
            with open(self.name_dic_path, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith("#"):
                        names.append(line)
        except OSError as e:
            self.logger.warning(f"无法读取 name_dic_for_user.txt: {e}")
            return []

        return names

    def load_parser(self):
        """第二步：加载解析器。

        根据文件类型返回对应的解析函数。
        支持 PDF、DOC、DOCX、TXT。

        Returns:
            解析函数
        """
        self.logger.info(f"加载解析器，文件类型: {self.file_ext}")

        if self.file_ext == ".pdf":
            try:
                import fitz  # PyMuPDF
                self.logger.info("成功加载 PDF 解析器 (PyMuPDF)")
                return self._parse_pdf
            except ImportError:
                raise ImportError("缺少 PyMuPDF 依赖，请运行: pip install PyMuPDF")

        elif self.file_ext == ".docx":
            try:
                from docx import Document
                self.logger.info("成功加载 DOCX 解析器 (python-docx)")
                return self._parse_docx
            except ImportError:
                raise ImportError("缺少 python-docx 依赖，请运行: pip install python-docx")

        elif self.file_ext == ".txt":
            self.logger.info("使用内置解析器处理 TXT")
            return self._parse_txt

        elif self.file_ext == ".doc":
            try:
                import textract
                self.logger.info("成功加载 DOC 解析器 (textract)")
                return self._parse_doc
            except ImportError:
                raise ImportError(
                    "缺少 textract 依赖，请运行: pip install textract。"
                    "注意：.doc 解析需要系统安装 antiword(Linux) 或相应后端(Windows)"
                )

        else:
            raise ValueError(f"没有可用的解析器: {self.file_ext}")

    def _parse_pdf(self, file_path: str) -> str:
        """PDF 文件解析实现。

        使用 PyMuPDF 逐页读取 PDF 内容，适合处理大文件（几十万字）。
        PyMuPDF 性能优于 pdfplumber，对中文支持更好。

        Args:
            file_path: PDF 文件路径

        Returns:
            提取的文本内容
        """
        import fitz  # PyMuPDF

        self.logger.info(f"开始解析 PDF: {file_path}")
        text_parts = []
        total_chars = 0

        # 打开 PDF 文件
        doc = fitz.open(file_path)
        page_count = len(doc)
        self.logger.info(f"PDF 共 {page_count} 页")

        try:
            for i in range(page_count):
                page = doc[i]
                # 提取文本，保留布局信息
                page_text = page.get_text("text") or ""
                if page_text.strip():
                    text_parts.append(page_text)
                    total_chars += len(page_text)

                # 每 10 页报告一次进度
                if (i + 1) % 10 == 0:
                    self.logger.info(f"已处理 {i + 1}/{page_count} 页，当前字符数: {total_chars}")

        finally:
            doc.close()

        full_text = "\n".join(text_parts)
        self.logger.info(f"PDF 解析完成，总字符数: {len(full_text)}")
        return full_text

    def _parse_docx(self, file_path: str) -> str:
        """DOCX 文件解析实现。

        使用 python-docx 提取段落和表格中的文本。

        Args:
            file_path: DOCX 文件路径

        Returns:
            提取的文本内容
        """
        from docx import Document

        self.logger.info(f"开始解析 DOCX: {file_path}")
        text_parts = []

        doc = Document(file_path)

        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" ".join(row_texts))

        full_text = "\n".join(text_parts)
        self.logger.info(f"DOCX 解析完成，总字符数: {len(full_text)}")
        return full_text

    def _parse_doc(self, file_path: str) -> str:
        """DOC 文件解析实现。

        使用 textract 提取 .doc (Office 97-2003) 格式的文本。
        需要系统安装 antiword(Linux) 或相应依赖(Windows)。

        Args:
            file_path: DOC 文件路径

        Returns:
            提取的文本内容
        """
        import textract

        self.logger.info(f"开始解析 DOC: {file_path}")

        try:
            # textract 返回 bytes，需解码
            raw = textract.process(file_path)
            # 尝试多种编码
            for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
                try:
                    full_text = raw.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                full_text = raw.decode("utf-8", errors="replace")
        except Exception as e:
            raise RuntimeError(f"DOC 解析失败: {e}") from e

        self.logger.info(f"DOC 解析完成，总字符数: {len(full_text)}")
        return full_text

    def _parse_txt(self, file_path: str) -> str:
        """TXT 文件解析实现。

        直接读取纯文本文件，支持多种编码自动检测。

        Args:
            file_path: TXT 文件路径

        Returns:
            文件文本内容
        """
        self.logger.info(f"开始读取 TXT: {file_path}")

        for encoding in ("utf-8", "utf-8-sig", "gbk", "gb2312", "big5", "latin-1"):
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    full_text = f.read()
                self.logger.info(f"TXT 读取完成，编码: {encoding}，总字符数: {len(full_text)}")
                return full_text
            except UnicodeDecodeError:
                continue

        raise RuntimeError(f"无法解码文件 {file_path}，尝试的编码均失败")

    def read_content(self, parser) -> str:
        """第三步：将文件内容读取为文本。

        使用指定的解析器读取文件内容。
        适合处理大文件（几十万字），采用流式/分块读取策略。

        Args:
            parser: 解析函数

        Returns:
            文件文本内容
        """
        self.logger.info("开始读取文件内容")
        content = parser(self.file_path)
        self.logger.info(f"内容读取完成，总长度: {len(content)} 字符")
        return content

    def process(self) -> str:
        """执行完整的准备流程。

        步骤：
        1. 确保 spaCy 模型已安装
        2. 校验文件
        3. 加载解析器
        4. 读取内容

        Returns:
            提取的文件文本内容

        Raises:
            各种校验或解析错误
        """
        # 第一步：确保 spaCy 模型（必需）
        self.ensure_spacy_models()

        # 第二步：加载用户自定义人名词典
        self.custom_names = self._load_custom_names()
        if self.custom_names:
            self.logger.info(f"已加载 {len(self.custom_names)} 个自定义人名: {self.name_dic_path}")

        # 第三步：校验
        self.validate()

        # 第四步：加载解析器
        parser = self.load_parser()

        # 第五步：读取内容
        content = self.read_content(parser)

        return content
